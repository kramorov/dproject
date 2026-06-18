# svg_converter/pdf_to_docx.py
"""
Извлечение форматированного текста из PDF и экспорт в DOCX.

Использует PyMuPDF для извлечения текста с информацией о шрифтах,
и python-docx для создания Word-документа с сохранением форматирования.

Pipeline:
  1. Открыть PDF через fitz
  2. Для каждой страницы: page.get_text("dict") → текстовые блоки
  3. Сортировка блоков в порядке чтения (y↓, x→)
  4. Группировка последовательных спанов с одинаковым шрифтом в runs
  5. Сборка DOCX: параграфы, bold/italic, размер шрифта
"""
import io
from typing import List, Dict, Any


def _fitz_font_to_docx(span: dict) -> dict:
    """
    Конвертировать информацию о шрифте из fitz-спана в параметры для docx.
    Возвращает {name, size_pt, bold, italic}.
    """
    font = span.get('font', 'Arial')
    size_pt = round(span.get('size', 11), 1)
    flags = span.get('flags', 0)

    # fitz flags: 2^0=superscript, 2^1=italic, 2^2=serif, 2^3=mono, 2^4=bold
    bold = bool(flags & 16)
    italic = bool(flags & 2)

    # Убираем технический суффикс из имени шрифта (fitz добавляет -Bold, -Italic и т.д.)
    clean_font = font.split('-')[0] if '-' in font else font
    if not clean_font or clean_font in ('Unknown', 'unknown'):
        clean_font = 'Arial'

    return {
        'name': clean_font,
        'size_pt': size_pt,
        'bold': bold,
        'italic': italic,
    }


def extract_text_blocks(pdf_bytes: bytes) -> List[List[dict]]:
    """
    Извлечь текстовые блоки из PDF, постранично.

    Каждая страница → список блоков.
    Каждый блок → dict с ключами: text, font, size, bold, italic, bbox.
    """
    try:
        import fitz
    except ImportError:
        raise RuntimeError('PyMuPDF (fitz) не установлен: pip install PyMuPDF')

    doc = fitz.open(stream=pdf_bytes, filetype='pdf')
    pages = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        # get_text("dict") возвращает словарь с блоками
        text_dict = page.get_text('dict')
        blocks = text_dict.get('blocks', [])
        page_spans = []

        for block in blocks:
            if block.get('type') != 0:  # только текстовые блоки (type=0)
                continue

            for line in block.get('lines', []):
                line_text = ''
                line_fonts = []
                line_bbox = None

                for span in line.get('spans', []):
                    text = span.get('text', '')
                    if not text.strip():
                        continue

                    font_info = _fitz_font_to_docx(span)
                    line_text += text + ' '
                    line_fonts.append(font_info)
                    if line_bbox is None:
                        b = span.get('bbox')
                        if b:
                            line_bbox = list(b)
                        else:
                            line_bbox = [0, 0, 0, 0]

                if line_text.strip() and line_fonts:
                    # Используем шрифт первого спана как основной для строки
                    primary_font = line_fonts[0]
                    page_spans.append({
                        'text': line_text.strip(),
                        'font': primary_font['name'],
                        'size': primary_font['size_pt'],
                        'bold': any(f['bold'] for f in line_fonts),
                        'italic': any(f['italic'] for f in line_fonts),
                        'bbox': line_bbox or [0, 0, 0, 0],
                        'page': page_num + 1,
                    })

        # Извлечение таблиц
        tables = _extract_tables_from_page(page, page_num)
        for t in tables:
            page_spans.append(t)

        # Сортировка: сначала по Y, потом по X (таблицы и текст вперемешку)
        page_spans.sort(key=lambda s: (round(s['bbox'][1], -1), s['bbox'][0]))
        pages.append(page_spans)

    doc.close()
    return pages


def _extract_tables_from_page(page, page_num: int) -> List[dict]:
    """Извлечь таблицы со страницы через page.find_tables()."""
    try:
        tabs = page.find_tables()
    except Exception:
        return []

    result = []
    for table in tabs:
        cells = []
        bbox = list(table.bbox)
        for row in table.extract():
            row_cells = []
            for cell in row:
                row_cells.append(str(cell).strip() if cell else '')
            cells.append(row_cells)
        if cells:
            result.append({
                'type': 'table',
                'cells': cells,
                'bbox': bbox,
                'page': page_num + 1,
            })
    return result


def spans_to_docx(pages: List[List[dict]]) -> bytes:
    """
    Собрать DOCX из списка текстовых спанов.

    Группирует последовательные спаны с одинаковым форматированием
    в параграфы и таблицы. Значительный разрыв по вертикали
    начинает новый параграф.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError('python-docx не установлен: pip install python-docx')

    doc = Document()

    # Настройка страницы
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)

    # Стиль по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for page_num, page_spans in enumerate(pages, 1):
        if not page_spans:
            continue

        # Заголовок страницы
        if len(pages) > 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'— Страница {page_num} —')
            run.font.size = Pt(9)
            run.italic = True

        current_para = None
        prev_y = None
        # Обрабатываем элементы в порядке сортировки (текст и таблицы)
        for item in page_spans:
            if item.get('type') == 'table':
                _add_table_to_docx(doc, item)
                current_para = None
                prev_y = None
                continue

            span = item
            y = span['bbox'][1]
            text = span['text']

            # Новый параграф: если вертикальный разрыв > 1.5× высоты строки
            # или если это первый span на странице
            new_para = (
                current_para is None
                or (prev_y is not None and y - prev_y > span['size'] * 1.5)
            )

            if new_para:
                current_para = doc.add_paragraph()
                current_para.paragraph_format.space_after = Pt(2)
                current_para.paragraph_format.space_before = Pt(0)

            # Добавляем run с форматированием
            run = current_para.add_run(text)
            run.font.name = span['font']
            run.font.size = Pt(span['size'])
            run.bold = span['bold']
            run.italic = span['italic']

            prev_y = y

    # Сохраняем в BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _add_table_to_docx(doc, table_item: dict):
    """Добавить таблицу python-docx в документ."""
    cells = table_item.get('cells', [])
    if not cells:
        return

    rows = len(cells)
    cols = max(len(row) for row in cells) if cells else 1
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')

    for i, row_data in enumerate(cells):
        for j, cell_text in enumerate(row_data):
            if j >= cols:
                break
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = 'Arial'

    # Отступ после таблицы
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)


def _ocr_page_to_text(page_image) -> str:
    """Распознать текст с изображения страницы через Tesseract (rus+eng)."""
    try:
        import pytesseract
    except ImportError:
        raise RuntimeError('pytesseract не установлен: pip install pytesseract')

    # Tesseract должен быть установлен отдельно: https://github.com/UB-Mannheim/tesseract/wiki
    text = pytesseract.image_to_string(page_image, lang='rus+eng')
    return text


def _pdf_to_docx_via_ocr(pdf_bytes: bytes) -> bytes:
    """PDF → DOCX через OCR (Tesseract). Для сканированных документов."""
    try:
        import fitz
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError as e:
        raise RuntimeError(f'Не хватает библиотеки: {e}')

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    pdf_doc = fitz.open(stream=pdf_bytes, filetype='pdf')

    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        # Рендерим страницу в высоком разрешении для OCR
        pix = page.get_pixmap(dpi=300)
        img = Image.frombytes('RGB', [pix.width, pix.height], pix.samples)

        text = _ocr_page_to_text(img)

        if page_num > 0:
            doc.add_page_break()

        for line in text.split('\n'):
            line = line.strip()
            if line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(10)
                run.font.name = 'Arial'
            else:
                doc.add_paragraph()  # пустая строка

    pdf_doc.close()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def pdf_to_docx(pdf_bytes: bytes, strip_images: bool = False, ocr: bool = False) -> bytes:
    """
    Полный pipeline: PDF → форматированный DOCX.

    Пробует pdf2docx (специализированная библиотека), при отсутствии —
    фолбек на fitz + python-docx.

    Args:
        pdf_bytes: содержимое PDF-файла
        strip_images: удалить растровые изображения перед конвертацией
                      (полезно для PDF со сканом + OCR-слоем)
        ocr: использовать Tesseract OCR вместо pdf2docx (для сканов)

    Returns:
        DOCX-файл как bytes
    """
    if ocr:
        return _pdf_to_docx_via_ocr(pdf_bytes)

    # ── Приоритет: pdf2docx (лучшее сохранение структуры) ──
    try:
        result = _pdf_to_docx_via_pdf2docx(pdf_bytes)
    except ImportError:
        result = None
    except Exception:
        result = None

    if result is None:
        pages = extract_text_blocks(pdf_bytes)
        result = spans_to_docx(pages)

    if strip_images:
        result = strip_images_from_docx(result)

    return result


def strip_images_from_docx(docx_bytes: bytes) -> bytes:
    """
    Удалить все изображения из DOCX — и ссылки в XML, и файлы картинок.

    DOCX — это ZIP-архив. Удаляем:
    1. Все <w:drawing> элементы из document.xml
    2. Все файлы из word/media/
    """
    import zipfile
    import io as _io

    buf = _io.BytesIO()
    with zipfile.ZipFile(_io.BytesIO(docx_bytes), 'r') as zin:
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                content = zin.read(item.filename)

                # Пропускаем папку с картинками
                if item.filename.startswith('word/media/'):
                    continue

                # Из document.xml удаляем все <w:drawing>...</w:drawing>
                if item.filename == 'word/document.xml':
                    import re
                    content = re.sub(
                        rb'<w:drawing[^>]*>.*?</w:drawing>',
                        b'',
                        content,
                        flags=re.DOTALL,
                    )

                zout.writestr(item, content)

    buf.seek(0)
    return buf.getvalue()


def _pdf_to_docx_via_pdf2docx(pdf_bytes: bytes) -> bytes:
    """Конвертация через pdf2docx (сохраняет таблицы, колонки, стили)."""
    from pdf2docx import Converter

    buf = io.BytesIO()

    # Сохраняем PDF во временный файл (pdf2docx требует путь на диске)
    import tempfile, os
    tmp_pdf = None
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            f.write(pdf_bytes)
            tmp_pdf = f.name

        # Конвертируем
        cv = Converter(tmp_pdf)
        cv.convert(buf)
        cv.close()

        buf.seek(0)
        return buf.getvalue()

    finally:
        if tmp_pdf and os.path.isfile(tmp_pdf):
            try:
                os.unlink(tmp_pdf)
            except OSError:
                pass