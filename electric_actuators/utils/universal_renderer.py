# electric_actuators/utils/universal_renderer.py pip install docxtpl
from jinja2 import Environment , FileSystemLoader , Template
from docxtpl import DocxTemplate
from pathlib import Path
import io
import logging
from datetime import datetime
from django.http import HttpResponse , JsonResponse

logger = logging.getLogger(__name__)
import markdown
from markdown.extensions import tables, extra


class UniversalTemplateRenderer:
    """
    Универсальный рендерер для HTML и Word документов.
    Использует один Jinja2 шаблон для всех форматов.
    """

    def __init__(self, template_dir=None):
        """Инициализация"""
        if template_dir is None:
            self.template_dir = Path(__file__).parent.parent / 'templates'
        else:
            self.template_dir = Path(template_dir)
        # Настройка Jinja2 для HTML
        self.jinja_env = Environment(
            loader=FileSystemLoader(self.template_dir),
            trim_blocks=True,
            lstrip_blocks=True,
            autoescape=False
        )
        self.ea_description_template_file_name = 'description_template.j2'
        # Регистрируем пользовательские фильтры
        self._register_filters()

    def _register_filters(self):
        """Регистрирует пользовательские фильтры Jinja2"""

        def format_currency(value, currency='₽'):
            """Форматирует валюту"""
            try:
                return f"{float(value):,.2f} {currency}".replace(',', ' ')
            except:
                return value

        def default_if_none(value, default='—'):
            """Заменяет None на значение по умолчанию"""
            return value if value is not None else default

        def format_bool(value, true_text='Да', false_text='Нет'):
            """Форматирует булево значение"""
            return true_text if value else false_text

        self.jinja_env.filters['currency'] = format_currency
        self.jinja_env.filters['default'] = default_if_none
        self.jinja_env.filters['bool'] = format_bool

    def prepare_data(self, raw_data):
        """
        Подготавливает данные из _generate_data_for_description для шаблона

        Args:
            raw_data: словарь из _generate_data_for_description()

        Returns:
            dict: подготовленные данные для шаблона
        """
        prepared = {
            'model': raw_data.get('model', {}),
            'code': raw_data.get('model', {}).get('name'),
            'brand': raw_data.get('brand', {}),
            'generated_at': datetime.now().strftime('%d.%m.%Y %H:%M'),
            'basic_params': [],
            'electrical_params': [],
            'mechanical_params': [],
        }

        # Определяем категории параметров
        electrical_keys = ['power_supply', 'motor_current_rated', 'motor_current_starting', 'motor_power']
        mechanical_keys = ['time_to_open', 'time_to_close', 'torque_min', 'torque_max',
                           'mounting_plate', 'stem_shape', 'stem_size', 'cable_glands_holes']

        # Единицы измерения
        units = {
            'power_supply': 'В',
            'motor_current_rated': 'А',
            'motor_current_starting': 'А',
            'motor_power': 'кВт',
            'time_to_open': 'с',
            'time_to_close': 'с',
            'torque_min': 'Нм',
            'torque_max': 'Нм',
        }

        # Обрабатываем все ключи из raw_data
        for key, value in raw_data.items():
            if key == 'model' or key == 'brand' or key == 'code':
                continue

            # Добавляем единицу измерения если есть
            if isinstance(value, dict):
                value_copy = value.copy()
                if key in units:
                    value_copy['unit'] = units[key]
                else:
                    value_copy['unit'] = None

                # Категоризируем
                if key in electrical_keys:
                    prepared['electrical_params'].append(value_copy)
                elif key in mechanical_keys:
                    prepared['mechanical_params'].append(value_copy)
                elif key.startswith('selected_') or key.endswith('_data'):
                    # Это составные опции, сохраняем как есть для детального вывода
                    prepared[key] = value_copy
                else:
                    prepared['basic_params'].append(value_copy)

        # Сортируем параметры по display_name
        for param_list in ['basic_params', 'electrical_params', 'mechanical_params']:
            prepared[param_list].sort(key=lambda x: x.get('display_name', ''))

        return prepared

    def render_html(self, template_name, data, convert_markdown=True):
        """
        Рендерит HTML из шаблона с опциональной конвертацией Markdown

        Args:
            template_name: имя файла шаблона
            data: подготовленные данные
            convert_markdown: конвертировать ли Markdown в HTML

        Returns:
            str: HTML строка
        """
        template = self.jinja_env.get_template(template_name)
        raw_html = template.render(**data)

        if convert_markdown:
            # Конвертируем Markdown в HTML
            md = markdown.Markdown(extensions=['tables', 'extra'])
            return md.convert(raw_html)

        return raw_html

    def render_docx(self, template_name, data):
        """
        Рендерит Word документ из того же шаблона
        """
        # Для Word используем docxtpl
        template_path = self.template_dir / template_name

        # Конвертируем .j2 в .docx если нужно
        docx_template_path = template_path.with_suffix('.docx')

        # Если нет .docx версии, создаем её из .j2
        if not docx_template_path.exists():
            self._create_docx_template(template_path, docx_template_path)

        doc = DocxTemplate(str(docx_template_path))

        # Подготавливаем данные для docxtpl
        docx_data = self._prepare_for_docx(data)

        # Для отладки
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"DOCX data keys: {docx_data.keys()}")

        doc.render(docx_data)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return output.getvalue()

    def _create_docx_template(self, jinja_path, docx_path):
        """Создает Word шаблон из Jinja2 шаблона с правильными переменными"""
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        doc = docx.Document()

        # Настройка стилей
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # ========== ЗАГОЛОВОК ==========
        title = doc.add_heading('ОПИСАНИЕ ЭЛЕКТРОПРИВОДА', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(16)
        title.runs[0].font.bold = True

        doc.add_paragraph()

        # ========== МОДЕЛЬ ==========
        doc.add_heading('Модель', level=1)

        p = doc.add_paragraph()
        p.add_run('{{ model_display_name }}: ').bold = True
        p.add_run('{{ model_name }}')

        p = doc.add_paragraph()
        p.add_run('Код: ').bold = True
        p.add_run('{{ code }}')

        p = doc.add_paragraph()
        p.add_run('{{ brand_display_name }}: ').bold = True
        p.add_run('{{ brand_value }}')

        doc.add_paragraph('---')

        # ========== ОСНОВНЫЕ ПАРАМЕТРЫ ==========
        doc.add_heading('Основные параметры', level=1)

        # Создаем таблицу
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        table.autofit = True

        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'
        hdr_cells[2].text = 'Тип'

        # Задаем ширину колонок
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Добавляем строку с циклом
        row_cells = table.add_row().cells
        row_cells[0].text = '{% for param in basic_params %}'
        row_cells[0].paragraphs[0].runs[0].font.italic = True
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        # Строка с данными
        row_cells = table.add_row().cells
        row_cells[0].text = '{{ param.display_name }}'
        row_cells[1].text = '{{ format_value(param.value, param.unit) }}'
        row_cells[2].text = '{% if param.is_default %}✓ Стандарт{% else %}⟳ Опция{% endif %}'

        # Закрываем цикл
        row_cells = table.add_row().cells
        row_cells[0].text = '{% endfor %}'
        row_cells[0].paragraphs[0].runs[0].font.italic = True
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        doc.add_paragraph()

        # ========== ЭЛЕКТРИЧЕСКИЕ ПАРАМЕТРЫ ==========
        doc.add_heading('Электрические параметры', level=1)

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'
        hdr_cells[2].text = 'Тип'

        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        row_cells = table.add_row().cells
        row_cells[0].text = '{% for param in electrical_params %}'
        row_cells[0].paragraphs[0].runs[0].font.italic = True
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        row_cells = table.add_row().cells
        row_cells[0].text = '{{ param.display_name }}'
        row_cells[1].text = '{{ format_value(param.value, param.unit) }}'
        row_cells[2].text = '{% if param.is_default %}✓ Стандарт{% else %}⟳ Опция{% endif %}'

        row_cells = table.add_row().cells
        row_cells[0].text = '{% endfor %}'
        row_cells[0].paragraphs[0].runs[0].font.italic = True
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        doc.add_paragraph()

        # ========== МЕХАНИЧЕСКИЕ ПАРАМЕТРЫ ==========
        doc.add_heading('Механические параметры', level=1)

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'
        hdr_cells[2].text = 'Тип'

        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        row_cells = table.add_row().cells
        row_cells[0].text = '{% for param in mechanical_params %}'
        row_cells[0].paragraphs[0].runs[0].font.italic = True
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        row_cells = table.add_row().cells
        row_cells[0].text = '{{ param.display_name }}'
        row_cells[1].text = '{{ format_value(param.value, param.unit) }}'
        row_cells[2].text = '{% if param.is_default %}✓ Стандарт{% else %}⟳ Опция{% endif %}'

        row_cells = table.add_row().cells
        row_cells[0].text = '{% endfor %}'
        row_cells[0].paragraphs[0].runs[0].font.italic = True
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        doc.add_paragraph()
        doc.add_paragraph('---')

        # ========== ДОПОЛНИТЕЛЬНЫЕ ОПЦИИ ==========
        doc.add_heading('Дополнительные опции', level=1)

        # IP защита
        doc.add_heading('IP защита', level=2)
        p = doc.add_paragraph('{% if ip_data %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        p = doc.add_paragraph()
        p.add_run('{{ ip_data.ip_option.display_name }}: ').bold = True
        p.add_run('{{ ip_data.ip_option.value }}')

        p = doc.add_paragraph()
        p.add_run('Степень защиты: ').bold = True
        p.add_run('IP{{ ip_data.ip_rank.value }}')

        p = doc.add_paragraph()
        p.add_run('Тип: ').bold = True
        p.add_run('{% if ip_data.is_default.value %}✓ Стандарт{% else %}⟳ Опция{% endif %}')

        p = doc.add_paragraph('{% endif %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        # Взрывозащита
        doc.add_heading('Взрывозащита', level=2)
        p = doc.add_paragraph('{% if exd_data %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        p = doc.add_paragraph()
        p.add_run('{{ exd_data.exd_option.display_name }}: ').bold = True
        p.add_run('{{ exd_data.exd_option.value }}')

        p = doc.add_paragraph()
        p.add_run('Тип: ').bold = True
        p.add_run('{% if exd_data.is_default.value %}✓ Стандарт{% else %}⟳ Опция{% endif %}')

        p = doc.add_paragraph('{% endif %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        # Температурный режим
        doc.add_heading('Температурный режим', level=2)
        p = doc.add_paragraph('{% if selected_temperature %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        p = doc.add_paragraph()
        p.add_run('Диапазон: ').bold = True
        p.add_run('{{ selected_temperature.temperature_range.value }}')

        p = doc.add_paragraph()
        p.add_run('Минимальная температура: ').bold = True
        p.add_run('{{ selected_temperature.work_temp_min.value }}°C')

        p = doc.add_paragraph()
        p.add_run('Максимальная температура: ').bold = True
        p.add_run('{{ selected_temperature.work_temp_max.value }}°C')

        p = doc.add_paragraph()
        p.add_run('Тип: ').bold = True
        p.add_run('{% if selected_temperature.is_default.value %}✓ Стандарт{% else %}⟳ Опция{% endif %}')

        p = doc.add_paragraph('{% endif %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        # Покрытие корпуса
        doc.add_heading('Покрытие корпуса', level=2)
        p = doc.add_paragraph('{% if selected_body_coating %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        p = doc.add_paragraph()
        p.add_run('{{ selected_body_coating.body_coating_option.display_name }}: ').bold = True
        p.add_run('{{ selected_body_coating.body_coating_option.value }}')

        p = doc.add_paragraph()
        p.add_run('Тип: ').bold = True
        p.add_run('{% if selected_body_coating.is_default.value %}✓ Стандарт{% else %}⟳ Опция{% endif %}')

        p = doc.add_paragraph('{% endif %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        # Ручной дублер
        doc.add_heading('Ручной дублер', level=2)
        p = doc.add_paragraph('{% if selected_hand_wheel %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        p = doc.add_paragraph()
        p.add_run('{{ selected_hand_wheel.hand_wheel_option.display_name }}: ').bold = True
        p.add_run('{{ selected_hand_wheel.hand_wheel_option.value }}')

        p = doc.add_paragraph()
        p.add_run('Тип: ').bold = True
        p.add_run('{% if selected_hand_wheel.is_default.value %}✓ Стандарт{% else %}⟳ Опция{% endif %}')

        p = doc.add_paragraph('{% endif %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        # Угол поворота
        doc.add_heading('Угол поворота', level=2)
        p = doc.add_paragraph('{% if selected_turn_angle_option %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        p = doc.add_paragraph()
        p.add_run('Угол: ').bold = True
        p.add_run('{{ selected_turn_angle_option.turn_angle.value }}°')

        p = doc.add_paragraph()
        p.add_run('Предел регулировки: ').bold = True
        p.add_run('±{{ selected_turn_angle_option.turn_angle_deviation_limit.value }}°')

        p = doc.add_paragraph()
        p.add_run('Тип: ').bold = True
        p.add_run('{% if selected_turn_angle_option.is_default.value %}✓ Стандарт{% else %}⟳ Опция{% endif %}')

        p = doc.add_paragraph('{% endif %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        # Блок управления
        doc.add_heading('Блок управления', level=2)
        p = doc.add_paragraph('{% if selected_control_unit_option %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        p = doc.add_paragraph()
        p.add_run('{{ selected_control_unit_option.control_unit.display_name }}: ').bold = True
        p.add_run('{{ selected_control_unit_option.control_unit.value }}')

        p = doc.add_paragraph()
        p.add_run('Тип: ').bold = True
        p.add_run('{% if selected_control_unit_option.is_default.value %}✓ Стандарт{% else %}⟳ Опция{% endif %}')

        p = doc.add_paragraph('{% endif %}')
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        doc.add_paragraph()
        doc.add_paragraph('---')

        # ========== ДАТА ГЕНЕРАЦИИ ==========
        p = doc.add_paragraph()
        p.add_run('*Сгенерировано: {{ generated_at }}*')
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = docx.shared.RGBColor(100, 100, 100)

        # Сохраняем документ
        doc.save(docx_path)

    def _prepare_for_docx(self, data):
        """Подготавливает данные для docxtpl, сохраняя структуру из prepare_data"""

        result = {}

        # 1. Сохраняем все ключи верхнего уровня как есть
        for key, value in data.items():
            result[key] = value

        # 2. Добавляем плоские версии для удобства (но не удаляем исходные)
        if 'model' in data and isinstance(data['model'], dict):
            result['model_display_name'] = data['model'].get('display_name', 'Модель')
            result['model_name'] = data['model'].get('name', '')

        if 'brand' in data and isinstance(data['brand'], dict):
            result['brand_display_name'] = data['brand'].get('display_name', 'Бренд')
            result['brand_value'] = data['brand'].get('value', '')

        # 3. Для сложных опций делаем плоские ключи (например ip_data_ip_option_value)
        for key in ['ip_data', 'exd_data', 'selected_temperature',
                    'selected_body_coating', 'selected_hand_wheel',
                    'selected_turn_angle_option', 'selected_control_unit_option']:
            if key in data and isinstance(data[key], dict):
                for subkey, subvalue in data[key].items():
                    if isinstance(subvalue, dict):
                        for subsubkey, subsubvalue in subvalue.items():
                            result[f"{key}_{subkey}_{subsubkey}"] = subsubvalue
                    else:
                        result[f"{key}_{subkey}"] = subvalue

        # 4. Для отладки
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"DOCX prepared keys: {list(result.keys())}")

        return result

    def response_html(self, request, instance_id):
        """HTTP ответ с HTML"""
        from electric_actuators.models.ea_actuator_selected import ElectricActuatorSelected

        instance = ElectricActuatorSelected.objects.get(id=instance_id)
        raw_data = instance._generate_data_for_description()
        html = self.render_html(self.ea_description_template_file_name, raw_data, convert_markdown=False)

        return HttpResponse(html)

    def response_json_html(self, request, instance_id):
        """JSON ответ с HTML для AJAX"""
        from electric_actuators.models.ea_actuator_selected import ElectricActuatorSelected

        instance = ElectricActuatorSelected.objects.get(id=instance_id)
        raw_data = instance._generate_data_for_description()

        html = self.render_html(self.ea_description_template_file_name, raw_data, convert_markdown=True)  # <-- Включено!

        return JsonResponse({
            'success': True,
            'html': html,
            'data': raw_data
        })

    def response_docx(self, request, instance_id):
        """HTTP ответ с Word документом"""
        from electric_actuators.models.ea_actuator_selected import ElectricActuatorSelected

        instance = ElectricActuatorSelected.objects.get(id=instance_id)
        raw_data = instance._generate_data_for_description()
        data = self.prepare_data(raw_data)

        docx_bytes = self.render_docx('description_template.j2', data)

        response = HttpResponse(
            docx_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="description_{instance_id}.docx"'

    # def render_docx(self, template_name, data):
    #     """
    #     Рендерит Word документ из .docx шаблона через docxtpl
    #     template_name: имя .docx файла (например 'description_template.docx')
    #     data: данные из _generate_data_for_description()
    #
    #     Returns:
    #         bytes: бинарные данные Word документа
    #     """
    #     from docxtpl import DocxTemplate
    #
    #     # Путь к .docx шаблону
    #     template_path = self.template_dir / template_name
    #
    #     if not template_path.exists():
    #         raise FileNotFoundError(f"Word шаблон не найден: {template_path}")
    #
    #     doc = DocxTemplate(str(template_path))
    #
    #     # Подготавливаем данные для docxtpl
    #     docx_data = self._prepare_for_docx(data)
    #
    #     # Добавляем дату генерации
    #     docx_data['generated_at'] = datetime.now().strftime('%d.%m.%Y %H:%M')
    #
    #     logger.info(f"DOCX rendering with keys: {list(docx_data.keys())}")
    #
    #     doc.render(docx_data)
    #
    #     output = io.BytesIO()
    #     doc.save(output)
    #     output.seek(0)
    #
    #     return output.getvalue()
    #
    # def docx_to_html(self, docx_bytes):
    #     """
    #     Конвертирует Word документ в HTML используя mammoth
    #
    #     Args:
    #         docx_bytes: бинарные данные Word документа
    #
    #     Returns:
    #         str: HTML строка
    #     """
    #     import mammoth
    #
    #     result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    #     return result.value
    #
    # def response_html(self, request, instance_id):
    #     """HTTP ответ с HTML (через Word -> HTML)"""
    #     from electric_actuators.models.ea_actuator_selected import ElectricActuatorSelected
    #
    #     instance = ElectricActuatorSelected.objects.get(id=instance_id)
    #     raw_data = instance._generate_data_for_description()
    #
    #     # 1. Сначала генерируем Word документ
    #     docx_bytes = self.render_docx('description_template.docx', raw_data)
    #
    #     # 2. Конвертируем Word в HTML
    #     html_content = self.docx_to_html(docx_bytes)
    #
    #     return HttpResponse(html_content)
    #
    # def response_json_html(self, request, instance_id):
    #     """JSON ответ с HTML для AJAX"""
    #     from electric_actuators.models.ea_actuator_selected import ElectricActuatorSelected
    #
    #     instance = ElectricActuatorSelected.objects.get(id=instance_id)
    #     raw_data = instance._generate_data_for_description()
    #
    #     # 1. Сначала генерируем Word документ
    #     docx_bytes = self.render_docx('description_template.docx', raw_data)
    #
    #     # 2. Конвертируем Word в HTML
    #     html_content = self.docx_to_html(docx_bytes)
    #
    #     return JsonResponse({
    #         'success': True,
    #         'html': html_content,
    #         'data': raw_data
    #     })
    #
    # def response_docx(self, request, instance_id):
    #     """HTTP ответ с Word документом (прямая загрузка)"""
    #     from electric_actuators.models.ea_actuator_selected import ElectricActuatorSelected
    #
    #     instance = ElectricActuatorSelected.objects.get(id=instance_id)
    #     raw_data = instance._generate_data_for_description()
    #
    #     docx_bytes = self.render_docx('description_template.docx', raw_data)
    #
    #     response = HttpResponse(
    #         docx_bytes,
    #         content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    #     )
    #     response['Content-Disposition'] = f'attachment; filename="description_{instance_id}.docx"'
    #
    #     return response